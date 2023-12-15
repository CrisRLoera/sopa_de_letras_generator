import random as rm
from docx import Document
from docx2pdf import convert

emptyCruz = []
wordsList = []

def addRowCol(matriz,row,col):
    for e in range(0,row):
        matriz.append([])
    for i in range(0,numRow):
        for j in range(0,col):
            matriz[i].append(1)
    return matriz

def isEmptyLR(matriz,ini,end,row):
    #print(ini,",",end,",",row)
    for i in range(ini,end):
        if(matriz[row][i] == 1):
            pass
        else:
            return True
    return False


def posWordLR(matriz,row,col,word):
    # Podria causar un bug despues xd
    ok = True
    while ok:
        iniRanNumRow = rm.randint(0,(row-1))
        iniRanNumCol = rm.randint(0,(col-1))
        if((iniRanNumCol+len(word))<len(matriz)):
            ok = isEmptyLR(matriz,iniRanNumCol,iniRanNumCol+len(word)-1,iniRanNumRow)
        else:
            pass

    if(((iniRanNumCol+len(word))>(col))):
        matriz = posWordLR(matriz,row,col,word)
        return matriz
    else:
        index = 0
        for i in range(iniRanNumCol,iniRanNumCol+len(word)):
            matriz[iniRanNumRow][i]=word[index]
            index +=1
        return matriz

def isEmptyUD(matriz,ini,end,col):
    #print(ini,",",end,",",row)
    for i in range(ini,end):
        if(matriz[i][col] == 1):
            pass
        else:
            return True
    return False


def posWordUD(matriz,row,col,word):
    # Podria causar un bug despues xd
    ok = True
    while ok:
        iniRanNumRow = rm.randint(0,(row-1))
        iniRanNumCol = rm.randint(0,(col-1))
        if((iniRanNumRow+len(word))<len(matriz[0])):
            ok = isEmptyUD(matriz,iniRanNumRow,iniRanNumRow+len(word)-1,iniRanNumCol)
        else:
            pass

    if(((iniRanNumCol+len(word))>(row))):
        matriz = posWordUD(matriz,row,col,word)
        return matriz
    else:
        index = 0
        for i in range(iniRanNumRow,iniRanNumRow+len(word)):
            matriz[i][iniRanNumCol]=word[index]
            index +=1
        return matriz


def posRandomWord(matriz,row,col):
    
    print("Ingresa el número de palabras")
    numWords = int(input())
    for i in range(0,numWords):
        print(f"Escribe la palabra número {i+1}:")
        i = input()
        wordsList.append(i)
    
    for i in range(0,numWords):
        selectFunction = rm.randint(1,2)
        if(selectFunction == 1):
            posWordLR(matriz,row,col,wordsList[i])
        elif(selectFunction == 2):
            posWordUD(matriz,row,col,wordsList[i])
        else:
            print("no funciono")

print("Ingresa el número de filas:")
numRow = int(input())
print("Ingresa el número de columnas:")
numCol = int(input())

emptyCruz = addRowCol(emptyCruz,numRow,numCol)
posRandomWord(emptyCruz,numRow,numCol)

abecedary = ['a','b','c','d','e','f','g','h','i','j','k','l','m','n','o','p','q','r','s','t','u','v','w','x','y','z']
print("")
for i in range(numRow):
    for j in range(numCol):
        if(emptyCruz[i][j] == 1):
            emptyCruz[i][j] = rm.choice(abecedary)


document = Document()
table = document.add_table(rows=numRow,cols=numCol)
for i in range(numRow):
    element = table.rows[i].cells
    for j in range(numCol):
        element[j].text = emptyCruz[i][j]
        
document.add_paragraph("")
list = document.add_paragraph("")
list.add_run("List").bold = True

text = document.add_paragraph("")
numElem = 0
for word in wordsList:
    numElem += 1
    text.add_run(f"{numElem}.-{word}    ")
document.save("sopa_de_letras.docx")

convert('sopa_de_letras.docx','sopa_de_letras.pdf')